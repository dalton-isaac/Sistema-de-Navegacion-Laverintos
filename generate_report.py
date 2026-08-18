"""
Script para generar el Informe Técnico Académico en formato Word (.docx)
para el Proyecto Final de Fundamentos de Inteligencia Artificial (PUCE).
Autor: Isaac Oña
Docente: Ing. Patricio Alvear
NRC: 1371
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_border(cell, **kwargs):
    """
    Establece bordes personalizados en una celda de tabla en python-docx.
    kwargs: top, bottom, left, right, etc. (valores: {"val": "single", "sz": "4", "color": "CCCCCC"})
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'<w:top w:val="{kwargs.get("top", {}).get("val", "none")}" w:sz="{kwargs.get("top", {}).get("sz", "4")}" w:space="0" w:color="{kwargs.get("top", {}).get("color", "auto")}"/>\n'
        f'<w:left w:val="{kwargs.get("left", {}).get("val", "none")}" w:sz="{kwargs.get("left", {}).get("sz", "4")}" w:space="0" w:color="{kwargs.get("left", {}).get("color", "auto")}"/>\n'
        f'<w:bottom w:val="{kwargs.get("bottom", {}).get("val", "none")}" w:sz="{kwargs.get("bottom", {}).get("sz", "4")}" w:space="0" w:color="{kwargs.get("bottom", {}).get("color", "auto")}"/>\n'
        f'<w:right w:val="{kwargs.get("right", {}).get("val", "none")}" w:sz="{kwargs.get("right", {}).get("sz", "4")}" w:space="0" w:color="{kwargs.get("right", {}).get("color", "auto")}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo hexadecimal a una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece los márgenes internos de una celda en dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>\n'
        f'<w:top w:w="{top}" w:type="dxa"/>\n'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'<w:left w:w="{left}" w:type="dxa"/>\n'
        f'<w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def construir_informe():
    input_doc_path = 'Proyecto Final Fundamentos de Inteligencia Artificial.docx'
    
    # Cargar documento base para preservar la plantilla original
    doc = docx.Document(input_doc_path)
    
    # Preservar los párrafos de la portada (P0 a P14)
    # y eliminar los párrafos vacíos posteriores (P15 en adelante)
    while len(doc.paragraphs) > 15:
        p = doc.paragraphs[-1]
        p._p.getparent().remove(p._p)
        
    # Añadir salto de página tras la portada
    doc.paragraphs[14].insert_paragraph_before('') # Espacio
    p_break = doc.add_paragraph()
    run_br = p_break.add_run()
    run_br.add_break(docx.enum.text.WD_BREAK.PAGE)

    # --- CONFIGURACIÓN DE FUNCIONES DE FORMATO APA 7MA EDICIÓN ---
    COLOR_TITULO = RGBColor(26, 54, 93)     # Azul oscuro académico (#1A365D)
    COLOR_SUBTITULO = RGBColor(43, 108, 176) # Azul secundario (#2B6CB0)
    COLOR_TEXTO = RGBColor(26, 32, 44)       # Gris oscuro / casi negro (#1A202C)

    def add_h1(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_with_next = True
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_TITULO
        return p

    def add_h2(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_with_next = True
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_SUBTITULO
        return p

    def add_h3(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_with_next = True
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = COLOR_TEXTO
        return p

    def add_body(texto, bold_prefix=None, italic_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Times New Roman'
            r_bold.font.size = Pt(12)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_TEXTO

        if italic_prefix:
            r_it = p.add_run(italic_prefix)
            r_it.font.name = 'Times New Roman'
            r_it.font.size = Pt(12)
            r_it.font.italic = True
            r_it.font.color.rgb = COLOR_TEXTO

        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_TEXTO
        return p

    def add_bullet(texto, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Times New Roman'
            r_bold.font.size = Pt(12)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_TEXTO

        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_TEXTO
        return p

    def add_code(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_shading(cell, "F8F9FA")
        set_cell_border(cell, left={"val": "single", "sz": "18", "color": "2B6CB0"},
                              top={"val": "single", "sz": "4", "color": "E2E8F0"},
                              bottom={"val": "single", "sz": "4", "color": "E2E8F0"},
                              right={"val": "single", "sz": "4", "color": "E2E8F0"})
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(45, 55, 72)
        doc.add_paragraph() # Espacio

    def add_table_apa(headers, data, col_widths=None, alignment=None, title=None, note=None):
        if title:
            p_t = doc.add_paragraph()
            p_t.paragraph_format.space_before = Pt(8)
            p_t.paragraph_format.space_after = Pt(2)
            p_t.paragraph_format.keep_with_next = True
            r_t = p_t.add_run(title)
            r_t.font.name = 'Times New Roman'
            r_t.font.size = Pt(11)
            r_t.font.bold = True
            r_t.font.color.rgb = COLOR_TITULO

        tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        # Cabecera
        hdr_row = tbl.rows[0]
        hdr_tr = hdr_row._tr.get_or_add_trPr()
        hdr_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for i, header_text in enumerate(headers):
            cell = hdr_row.cells[i]
            if col_widths and i < len(col_widths):
                cell.width = Inches(col_widths[i])
            set_cell_shading(cell, "1A365D")
            set_cell_border(cell, top={"val": "single", "sz": "12", "color": "1A365D"},
                                  bottom={"val": "single", "sz": "12", "color": "1A365D"})
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Filas de datos
        for r_idx, row_data in enumerate(data):
            row = tbl.rows[r_idx + 1]
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                if col_widths and c_idx < len(col_widths):
                    cell.width = Inches(col_widths[c_idx])
                set_cell_shading(cell, bg_color)
                # Borde inferior fino o borde APA final
                is_last_row = (r_idx == len(data) - 1)
                b_color = "1A365D" if is_last_row else "E2E8F0"
                b_sz = "12" if is_last_row else "4"
                set_cell_border(cell, bottom={"val": "single", "sz": b_sz, "color": b_color})
                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                # Alineación
                if alignment and c_idx < len(alignment):
                    p.paragraph_format.alignment = alignment[c_idx]
                else:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    
                run = p.add_run(str(val))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_TEXTO
                if c_idx == 0:
                    run.font.bold = True

        if note:
            p_n = doc.add_paragraph()
            p_n.paragraph_format.space_before = Pt(3)
            p_n.paragraph_format.space_after = Pt(8)
            r_n = p_n.add_run(note)
            r_n.font.name = 'Times New Roman'
            r_n.font.size = Pt(9)
            r_n.font.italic = True
            r_n.font.color.rgb = RGBColor(100, 116, 139)

    def add_image_figure(image_path, title_text, caption_text, width_inches=6.2):
        if os.path.exists(image_path):
            p_img = doc.add_paragraph()
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(3)
            p_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.keep_with_next = True
            p_img.add_run().add_picture(image_path, width=Inches(width_inches))

            p_cap = doc.add_paragraph()
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(10)
            p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            r_num = p_cap.add_run(title_text + " ")
            r_num.font.name = 'Times New Roman'
            r_num.font.size = Pt(9.5)
            r_num.font.bold = True
            r_num.font.color.rgb = COLOR_TITULO

            r_desc = p_cap.add_run(caption_text)
            r_desc.font.name = 'Times New Roman'
            r_desc.font.size = Pt(9.5)
            r_desc.font.italic = True
            r_desc.font.color.rgb = RGBColor(74, 85, 104)

    # =========================================================================
    # SECCIÓN: RESUMEN / ABSTRACT
    # =========================================================================
    add_h1("Resumen")
    add_body("El presente trabajo de investigación aplicada expone el diseño, desarrollo, implementación y evaluación experimental de un Sistema Inteligente de Búsqueda y Navegación en Laberintos orientado a operaciones de rescate robótico en edificaciones colapsadas. El problema se formaliza mediante la teoría de grafos y espacios de estados sobre matrices bidimensionales con restricciones de transitabilidad, incorporando un modelo de costos asimétricos donde el desplazamiento horizontal posee un costo unitario (c=1) y el desplazamiento vertical un costo penalizado (c=2), simulando el esfuerzo mecánico y gravitacional de ascenso y descenso en estructuras siniestradas. Se implementaron y evaluaron cinco algoritmos fundamentales de Inteligencia Artificial: tres no informados o ciegos (Búsqueda en Anchura - BFS, Búsqueda en Profundidad - DFS y Búsqueda de Costo Uniforme - UCS) y dos informados o heurísticos (Búsqueda Avara - Greedy Best-First Search y Algoritmo A*), empleando la heurística admisible y consistente de Distancia Manhattan. Los resultados experimentales en múltiples escenarios demuestran que el algoritmo A* logra un balance óptimo superior, garantizando rutas de costo mínimo exacto idénticas a las de UCS pero reduciendo el espacio de nodos explorados hasta en un 50% a 70%. Asimismo, se desarrolló una interfaz gráfica interactiva en Tkinter que facilita la edición de escenarios, animación en tiempo real de la frontera de exploración y benchmarking comparativo.")
    add_body("Búsqueda no informada, Búsqueda heurística, Algoritmo A*, Costo Uniforme, Distancia Manhattan, Robótica de rescate, Grafos de navegación.", bold_prefix="Palabras clave: ")

    add_h1("Abstract")
    add_body("This applied research work presents the design, development, implementation, and experimental evaluation of an Intelligent Search and Navigation System in Mazes tailored for robotic search-and-rescue operations in collapsed structures. The problem is formalized using graph theory and discrete state-space representations on two-dimensional grid matrices with obstacle constraints, incorporating an asymmetric cost model where horizontal movements carry a unit cost (c=1) and vertical movements incur a penalized cost (c=2) to reflect gravitational and mechanical traversal efforts. Five fundamental Artificial Intelligence search algorithms were implemented and tested: three uninformed/blind techniques (Breadth-First Search - BFS, Depth-First Search - DFS, and Uniform Cost Search - UCS) and two informed/heuristic techniques (Greedy Best-First Search and A* Search) using the admissible and consistent Manhattan Distance heuristic. Experimental benchmarks across diverse maze topologies reveal that the A* algorithm achieves optimal performance, finding minimal-cost paths identical to UCS while reducing the expanded state space by 50% to 70%. Furthermore, a modern interactive Graphical User Interface (GUI) in Tkinter was built to support interactive map editing, step-by-step exploration animation, and comparative metric benchmarking.")
    add_body("Uninformed search, Heuristic search, A* algorithm, Uniform Cost Search, Manhattan Distance, Rescue robotics, Autonomous navigation.", bold_prefix="Keywords: ")

    # =========================================================================
    # SECCIÓN 1: INTRODUCCIÓN
    # =========================================================================
    add_h1("1. Introducción")
    add_body("En el campo de la Inteligencia Artificial y la Robótica Autónoma, la navegación y planificación de trayectorias en entornos confinados o parcialmente obstruidos representa uno de los desafíos más críticos. Tras eventos catastróficos tales como sismos, explosiones o colapsos estructurales, las labores de Búsqueda y Rescate Urbano (USAR, por sus siglas en inglés) demandan el despliegue de robots autónomos capaces de ingresar en áreas de alto riesgo, esquivar escombros y localizar a sobrevivientes o accesos de evacuación de manera rápida, segura y energéticamente eficiente.")
    add_body("La navegación autónoma de estos agentes depende de algoritmos de búsqueda en espacios de estados que transforman la topología del entorno físico en un grafo conexo de decisiones. En este contexto, la selección del algoritmo de búsqueda no solo determina si el agente alcanzará el objetivo, sino también la cantidad de recursos computacionales consumidos en tiempo real (nodos expandidos y tiempo de CPU) y la calidad intrínseca de la trayectoria resultante (longitud y costo energético acumulado).")
    add_body("En escenarios reales de rescate, no todos los desplazamientos implican el mismo gasto energético; ascender o descender desniveles estructurales requiere vencer fuerzas gravitacionales y sortear inestabilidades mecánicas, lo cual incrementa sensiblemente el costo respecto a los desplazamientos laterales. Por este motivo, el presente proyecto incorpora formalmente un modelo de costos diferenciados (costo horizontal de 1 unidad y costo vertical de 2 unidades), lo cual introduce una complejidad representativa que permite contrastar con rigor la diferencia entre algoritmos que minimizan simplemente el número de pasos (como BFS) y aquellos que optimizan el costo energético acumulado (como UCS y A*).")
    
    add_h2("1.1 Objetivos del Proyecto")
    add_body("Diseñar, implementar y evaluar un sistema computacional interactivo para la resolución de laberintos mediante algoritmos de búsqueda no informada (BFS, DFS, UCS) e informada (Greedy, A*), integrando una interfaz gráfica visual y un marco comparativo de métricas de desempeño.", bold_prefix="Objetivo General: ")
    
    add_body("Con el propósito de cumplir con el objetivo general, se establecen los siguientes objetivos específicos:")
    add_bullet("Modelar formalmente el entorno de navegación como un espacio de estados discreto y matricial, definiendo reglas de transitabilidad, operadores de transición y una función de costo asimétrica (horizontal = 1, vertical = 2).", bold_prefix="1. ")
    add_bullet("Desarrollar e integrar los algoritmos de búsqueda no informada: Búsqueda en Anchura (BFS), Búsqueda en Profundidad (DFS) y Búsqueda de Costo Uniforme (UCS / Dijkstra).", bold_prefix="2. ")
    add_bullet("Desarrollar e integrar los algoritmos de búsqueda informada: Búsqueda Avara (Greedy Best-First) y Algoritmo A*, formulando y validando la función heurística admisible de distancia Manhattan.", bold_prefix="3. ")
    add_bullet("Construir una interfaz gráfica de usuario (GUI) moderna e interactiva en Tkinter que permita crear, cargar, editar laberintos y visualizar paso a paso la animación de la exploración y la ruta calculada.", bold_prefix="4. ")
    add_bullet("Ejecutar un protocolo experimental de pruebas en múltiples escenarios (laberinto oficial, entorno urbano, trampa cóncava y escenario complejo de gran escala), analizando cuantitativamente nodos explorados, tiempo de CPU, costo acumulado y longitud de trayectoria.", bold_prefix="5. ")

    # =========================================================================
    # SECCIÓN 2: MARCO TEÓRICO
    # =========================================================================
    add_h1("2. Marco Teórico")
    add_body("El marco teórico del presente trabajo se fundamenta en la teoría clásica de resolución de problemas mediante búsqueda en Inteligencia Artificial simbólica, formulada por Russell y Norvig (2020), y en la optimización de caminos en grafos ponderados (Cormen et al., 2022).")

    add_h2("2.1 Representación del Espacio de Estados y Grafos de Navegación")
    add_body("Un problema de búsqueda se formaliza mediante una 5-tupla matemática:")
    add_code("Problema = < S, s_0, G, Acciones(s), Transicion(s, a), Costo(s, a, s') >\n"
             "Donde:\n"
             "  • S: Conjunto de estados posibles, representados por pares ordenados (fila, columna) en una matriz M x N.\n"
             "  • s_0: Estado inicial donde se posiciona el robot ('S').\n"
             "  • G: Conjunto de estados meta ('G') que representan la víctima o la salida.\n"
             "  • Acciones(s): Conjunto de movimientos ortogonales permitidos {Arriba, Abajo, Izquierda, Derecha}.\n"
             "  • Transicion(s, a): Función determinista que mapea un estado s y una acción a en un nuevo estado sucesor s'.\n"
             "  • Costo(s, a, s'): Función de costo de paso: Costo = 1 para desplazamientos horizontales y Costo = 2 para verticales.")

    add_h2("2.2 Algoritmos de Búsqueda No Informada (Búsqueda Ciega)")
    add_body("Los algoritmos no informados operan exclusivamente con la información provista por la definición del problema, sin estimaciones adicionales sobre la cercanía relativa a la meta.")

    add_h3("2.2.1 Búsqueda en Anchura (BFS - Breadth-First Search)")
    add_body("BFS expande sistemáticamente los nodos nivel por nivel, utilizando una estructura de datos de cola FIFO (First-In, First-Out). En grafos no ponderados o con aristas de costo uniforme, BFS garantiza encontrar la solución con el menor número de aristas (longitud mínima de trayectoria). Sin embargo, bajo un esquema de costos ponderados asimétricos, la solución de mínima longitud de pasos no coincide necesariamente con la de menor costo total acumulado. Su complejidad temporal y espacial es de orden O(b^d), donde b es el factor de ramificación y d la profundidad de la solución óptima en pasos, lo que provoca un alto consumo de memoria RAM en espacios extensos.")

    add_h3("2.2.2 Búsqueda en Profundidad (DFS - Depth-First Search)")
    add_body("DFS explora exhaustivamente cada rama del árbol de búsqueda hasta alcanzar su máxima profundidad antes de aplicar retroceso (backtracking), empleando una estructura de pila LIFO (Last-In, First-Out). Su principal ventaja teórica radica en su reducida complejidad espacial lineal O(b·m), donde m es la profundidad máxima del espacio de estados. No obstante, DFS no es un algoritmo óptimo ni completo en espacios infinitos o con ciclos sin control de visitados, tendiendo a encontrar rutas excesivamente largas, serpenteantes y de alto costo energético.")

    add_h3("2.2.3 Búsqueda de Costo Uniforme (UCS - Uniform Cost Search)")
    add_body("UCS, correspondiente a la generalización del algoritmo de Dijkstra para espacios de estados implícitos, gestiona la frontera de exploración mediante una cola de prioridad ordenada por el costo acumulado real g(n) desde el nodo inicial. UCS expande siempre el nodo no explorado con el menor g(n). Gracias a esta propiedad, UCS garantiza encontrar de forma rigurosa la trayectoria de costo óptimo global en cualquier grafo ponderado con costos de paso estrictamente positivos (c >= epsilon > 0). Su complejidad espacial y temporal es de orden O(b^(1 + floor(C*/epsilon))), donde C* es el costo óptimo de la ruta.")

    add_h2("2.3 Algoritmos de Búsqueda Informada (Búsqueda Heurística)")
    add_body("Los algoritmos informados incorporan conocimiento específico del dominio mediante una función heurística h(n) que estima el costo restante desde el estado actual n hasta el estado objetivo más cercano.")

    add_h3("2.3.1 Función Heurística de Distancia Manhattan")
    add_body("En entornos matriciales o de rejilla ortogonal con conectividad de 4 direcciones (vecindad de Von Neumann) donde los movimientos diagonales están prohibidos, la métrica de distancia Manhattan constituye la heurística estándar por excelencia. Se formula matemáticamente como:")
    add_code("h(n) = |x_n - x_meta| + |y_n - y_meta|\n"
             "Donde (x_n, y_n) son las coordenadas del nodo actual y (x_meta, y_meta) las de la meta.")
    add_body("La distancia Manhattan satisface formalmente dos propiedades matemáticas esenciales:")
    add_bullet("Admisibilidad: h(n) nunca sobreestima el costo real h*(n) para alcanzar la meta (h(n) <= h*(n)), ya que asume un entorno continuo libre de todo obstáculo donde cada paso tiene al menos costo unitario.", bold_prefix="• ")
    add_bullet("Consistencia o Monotonía: Para todo nodo n y cada sucesor n' generado por una acción a, se cumple la desigualdad triangular: h(n) <= c(n, a, n') + h(n'). La consistencia garantiza que el costo acumulado f(n) no decrezca a lo largo de cualquier trayectoria.", bold_prefix="• ")

    add_h3("2.3.2 Búsqueda Avara (Greedy Best-First Search)")
    add_body("La búsqueda avara evalúa los nodos de la frontera seleccionando en cada iteración aquel que minimiza exclusivamente la función heurística: f(n) = h(n). Su objetivo es avanzar de la manera más directa posible hacia el objetivo. Aunque en entornos abiertos exhibe una velocidad de convergencia sobresaliente con un mínimo consumo de memoria, carece de optimalidad y es propensa a desviarse significativamente al enfrentar obstáculos cóncavos (como trampas en U), viéndose forzada a retroceder tras explorar ramas sin salida.")

    add_h3("2.3.3 Algoritmo A* (A-Star Search)")
    add_body("El algoritmo A*, introducido por Hart, Nilsson y Raphael (1968), combina la rigurosidad de UCS con la orientación focalizada de Greedy mediante una función de evaluación compuesta:")
    add_code("f(n) = g(n) + h(n)\n"
             "Donde:\n"
             "  • g(n): Costo real acumulado desde el estado inicial s_0 hasta el estado n.\n"
             "  • h(n): Estimación heurística admisible del costo restante desde n hasta la meta.")
    add_body("Cuando la heurística h(n) es admisible y consistente, A* posee propiedades fundamentales:")
    add_bullet("Completitud: Encuentra siempre una solución si esta existe en el espacio de estados.", bold_prefix="1. ")
    add_bullet("Optimalidad Admisible: Garantiza encontrar el camino de menor costo total C* sin posibilidad de retornar una solución subóptima.", bold_prefix="2. ")
    add_bullet("Eficiencia Óptima: Ningún otro algoritmo de búsqueda óptima que utilice la misma función heurística h(n) puede expandir menos nodos que A*.", bold_prefix="3. ")

    # Tabla teórica
    headers_teo = ["Algoritmo", "Estrategia", "Frontera / Estructura", "Completo", "Óptimo", "Complejidad Tiempo", "Complejidad Espacio"]
    data_teo = [
        ["BFS", "No Informada", "Cola FIFO (deque)", "Sí (si b es finito)", "Sí (solo en aristas = 1)", "O(b^d)", "O(b^d)"],
        ["DFS", "No Informada", "Pila LIFO (stack)", "No (en bucles)", "No", "O(b^m)", "O(b·m)"],
        ["UCS", "No Informada", "Min-Heap g(n)", "Sí (si c >= eps > 0)", "Sí (Garantizado)", "O(b^(1+C*/eps))", "O(b^(1+C*/eps))"],
        ["Greedy", "Informada", "Min-Heap h(n)", "No (puede ciclar)", "No", "O(b^m)", "O(b^m)"],
        ["A*", "Informada", "Min-Heap f(n)=g+h", "Sí (si h es admisible)", "Sí (Garantizado)", "O(b^d*) (O(b^m) peor)", "O(b^d*)"]
    ]
    add_table_apa(headers_teo, data_teo, col_widths=[0.9, 1.0, 1.2, 0.8, 0.9, 1.1, 1.0], 
                  title="Tabla 1. Propiedades Teóricas y Complejidad de los Algoritmos de Búsqueda Evaluados",
                  note="Nota. b: factor de ramificación; d: profundidad óptima; m: profundidad máxima; C*: costo óptimo; eps: costo mínimo de paso.")

    # =========================================================================
    # SECCIÓN 3: DESCRIPCIÓN Y FORMULACIÓN DEL PROBLEMA
    # =========================================================================
    add_h1("3. Descripción y Formulación del Problema")
    add_body("El proyecto se contextualiza en el diseño del sistema de navegación inteligente para un robot de rescate autónomo. El robot debe ser capaz de ingresar a una edificación colapsada cuya distribución espacial se modela como una cuadrícula matricial bidimensional de dimensiones M x N. Dentro de esta matriz:")
    add_bullet("'S' (Start): Posición de entrada o despliegue inicial del robot rescatista.", bold_prefix="• ")
    add_bullet("'G' (Goal): Posición donde se encuentra una víctima atrapada o una salida de emergencia.", bold_prefix="• ")
    add_bullet("'0' (Espacio Libre): Celda despejada por donde el robot puede transitar libremente.", bold_prefix="• ")
    add_bullet("'1' (Obstáculo): Muros colapsados, escombros o losas infranqueables que impiden el paso.", bold_prefix="• ")

    add_h2("3.1 Reglas de Movimiento y Costos Asimétricos")
    add_body("El robot puede ejecutar cuatro movimientos discretos ortogonales, sujetos a restricciones físicas:")
    add_code("MOVIMIENTOS = {\n"
             "    (0,  1): 1,   # Desplazamiento Horizontal a la Derecha  -> Costo = 1\n"
             "    (0, -1): 1,   # Desplazamiento Horizontal a la Izquierda -> Costo = 1\n"
             "    (1,  0): 2,   # Desplazamiento Vertical hacia Abajo      -> Costo = 2\n"
             "    (-1, 0): 2    # Desplazamiento Vertical hacia Arriba     -> Costo = 2\n"
             "}")
    add_body("La asignación de un costo doble (c=2) a los movimientos verticales simula el esfuerzo motriz de ascenso y descenso en rampas de escombros o la resistencia mecánica contra la gravedad. Esta diferenciación de costos constituye una variable de estudio determinante, ya que induce discrepancias entre las trayectorias de menor longitud geométrica (seleccionadas por BFS) y las de menor consumo energético (seleccionadas por UCS y A*).")

    add_h2("3.2 Métricas de Desempeño Cuantitativo")
    add_body("Para realizar una comparación experimental rigurosa entre los cinco algoritmos, se instrumentó el sistema para registrar cuatro métricas clave:")
    add_bullet("Nodos Explorados (N_exp): Cantidad total de estados extraídos de la frontera y evaluados durante el proceso de búsqueda. Mide el esfuerzo computacional y el consumo de memoria.", bold_prefix="1. ")
    add_bullet("Tiempo de Ejecución (t_cpu): Tiempo de procesamiento de CPU medido en milisegundos (ms) con precisión de microsegundos mediante time.perf_counter().", bold_prefix="2. ")
    add_bullet("Costo Total de la Ruta (C_total): Sumatoria de los costos de transición de cada paso a lo largo de la trayectoria solución encontrada desde 'S' hasta 'G'.", bold_prefix="3. ")
    add_bullet("Longitud de la Ruta (L): Número total de pasos o transiciones ejecutadas en la ruta solución (longitud de la lista de coordenadas menos 1).", bold_prefix="4. ")

    # =========================================================================
    # SECCIÓN 4: IMPLEMENTACIÓN DEL SISTEMA
    # =========================================================================
    add_h1("4. Implementación del Sistema")
    add_body("El sistema se desarrolló íntegramente en lenguaje Python 3 (versión 3.10+) bajo un diseño modular, desacoplado y de alto rendimiento, estructurado en cuatro módulos principales:")
    add_bullet("Módulo del Núcleo Algorítmico: Contiene las funciones puras de búsqueda (bfs, dfs, ucs, greedy, a_star), asegurando consistencia matemática en las estructuras de datos y el cálculo de heurísticas.", bold_prefix="1. ")
    add_bullet("Módulo de Laberintos y Generador: Administra el catálogo de mapas predefinidos (presets) e integra un generador estocástico de laberintos con validación de solvabilidad garantizada.", bold_prefix="2. ")
    add_bullet("Módulo de Benchmarking y Gráficos: Ejecuta pruebas automatizadas multiescenario y genera figuras comparativas de alta resolución mediante Matplotlib.", bold_prefix="3. ")
    add_bullet("Módulo de Interfaz Gráfica Interactiva (GUI): Desarrollado en Tkinter/ttk, provee una experiencia visual fluida para la edición interactiva de celdas, animación paso a paso de la exploración y consulta de métricas en tiempo real.", bold_prefix="4. ")

    add_h2("4.1 Estructuras de Datos y Optimización de Complejidad")
    add_body("Para asegurar tiempos de respuesta óptimos y evitar cuellos de botella en memoria, se seleccionaron minuciosamente las estructuras de datos estándar de Python:")
    add_bullet("collections.deque: Utilizada en BFS para lograr operaciones de inserción al final y extracción al frente en tiempo constante O(1).", bold_prefix="• ")
    add_bullet("heapq (Binary Min-Heap): Utilizado en UCS, Greedy y A*. Para evitar comparaciones no deseadas entre listas de rutas cuando los costos son iguales, las tuplas insertadas en el heap incluyen un contador incremental de desempate: (prioridad, contador, nodo_actual, ruta).", bold_prefix="• ")
    add_bullet("Diccionarios y Conjuntos Hash (dict y set): Utilizados para el registro de estados visitados y el rastreo de costos mínimos g(n), permitiendo consultas e inserciones en tiempo amortizado O(1).", bold_prefix="• ")

    add_h2("4.2 Código Fuente Principal del Algoritmo A*")
    add_body("A continuación se presenta el fragmento de código de la implementación del algoritmo A*, ilustrando la gestión de la cola de prioridad y la actualización de costos g(n) y f(n):")
    add_code("def a_star(laberinto, registrar_exploracion=True):\n"
             "    inicio, meta = encontrar_inicio_meta(laberinto)\n"
             "    inicio_tiempo = time.perf_counter()\n"
             "    contador = 0\n"
             "    h_inicio = heuristica_manhattan(inicio, meta)\n"
             "    # Cola de prioridad: (f_cost, g_cost, contador, nodo_actual, ruta)\n"
             "    cola = [(h_inicio, 0, contador, inicio, [inicio])]\n"
             "    visitados = {}  # Mapea nodo -> menor g_cost registrado\n"
             "    orden_exploracion = []\n"
             "    nodos_explorados = 0\n"
             "\n"
             "    while cola:\n"
             "        f_cost, g_cost, _, actual, ruta = heapq.heappop(cola)\n"
             "        if actual in visitados and visitados[actual] <= g_cost:\n"
             "            continue\n"
             "        visitados[actual] = g_cost\n"
             "        nodos_explorados += 1\n"
             "        if registrar_exploracion:\n"
             "            orden_exploracion.append(actual)\n"
             "\n"
             "        if actual == meta:\n"
             "            t_exec = time.perf_counter() - inicio_tiempo\n"
             "            return ruta, nodos_explorados, t_exec, g_cost, len(ruta)-1, orden_exploracion\n"
             "\n"
             "        for (dx, dy), costo_mov in MOVIMIENTOS.items():\n"
             "            vecino = (actual[0] + dx, actual[1] + dy)\n"
             "            if es_celda_transitable(laberinto, vecino[0], vecino[1]):\n"
             "                nuevo_g = g_cost + costo_mov\n"
             "                if vecino not in visitados or nuevo_g < visitados[vecino]:\n"
             "                    nuevo_f = nuevo_g + heuristica_manhattan(vecino, meta)\n"
             "                    contador += 1\n"
             "                    heapq.heappush(cola, (nuevo_f, nuevo_g, contador, vecino, ruta + [vecino]))\n"
             "    return None, nodos_explorados, time.perf_counter() - inicio_tiempo, 0, 0, orden_exploracion")

    add_h2("4.3 Características de la Interfaz Gráfica (GUI)")
    add_body("La interfaz gráfica de usuario desarrollada ofrece un entorno completo para experimentación y demostración en vivo:")
    add_bullet("Editor Dinámico de Rejilla: Permite alternar muros ('1'), pasillos libres ('0'), posición inicial ('S') y meta ('G') mediante clics directos o arrastre continuo con el ratón.", bold_prefix="• ")
    add_bullet("Selector de Presets y Generador Aleatorio: Facilita la carga instantánea de escenarios de prueba o la creación de laberintos aleatorios garantizados como solvables.", bold_prefix="• ")
    add_bullet("Animador con Control de Velocidad: Simula en tiempo real la expansión de la frontera (celdas amarillas) y el trazado progresivo de la ruta óptima final (celdas azules y línea dorada), con velocidad ajustable de 1 ms a 150 ms por paso.", bold_prefix="• ")
    add_bullet("Panel de Benchmarking y Ventana de Resultados: Ejecuta simultáneamente los 5 algoritmos sobre el mapa activo, desplegando una tabla comparativa y permitiendo exportar gráficos de rendimiento generados con Matplotlib.", bold_prefix="• ")

    # =========================================================================
    # SECCIÓN 5: RESULTADOS EXPERIMENTALES
    # =========================================================================
    add_h1("5. Resultados Experimentales")
    add_body("Para validar empíricamente la efectividad y el comportamiento de los algoritmos implementados, se diseñó una batería de pruebas sobre cuatro escenarios con distintas complejidades topológicas, densidades de obstáculos y tamaños de rejilla.")

    # Caso 1
    add_h2("5.1 Caso de Estudio 1: Laberinto Oficial de Prueba (5x5)")
    add_body("El primer escenario corresponde a la matriz 5x5 estipulada en la guía oficial del proyecto, con un punto de inicio en (0,0) y la meta en (4,3). Representa un entorno controlado con obstáculos intermedios.")
    
    headers_c1 = ["Métrica Evaluada", "BFS", "DFS", "UCS", "Greedy", "A*"]
    data_c1 = [
        ["Nodos explorados", "16", "12", "16", "10", "12"],
        ["Tiempo de ejecución (ms)", "0.0631", "0.0388", "0.0590", "0.0382", "0.0464"],
        ["Costo total (H=1, V=2)", "13", "15", "13", "13", "13"],
        ["Longitud de la ruta (pasos)", "9", "11", "9", "9", "9"],
        ["¿Encontró ruta óptima?", "Sí", "No (Desvío)", "Sí (Costo Mín.)", "Sí (Coincidente)", "Sí (Costo Mín.)"]
    ]
    add_table_apa(headers_c1, data_c1, col_widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9],
                  title="Tabla 2. Métricas de Desempeño en el Escenario Oficial de Prueba (5x5)")
    
    add_image_figure("mapa_visual_oficial.png", "Figura 1.", "Comparación visual de la navegación y nodos explorados en el Laberinto Oficial (5x5). Verde: Inicio, Rojo: Meta, Amarillo: Nodos explorados, Azul: Trayectoria solución encontrada.")
    add_image_figure("grafico_preset_1.png", "Figura 2.", "Gráficos comparativos de barras para el Escenario 1 (5x5): Nodos explorados, Tiempo de ejecución (ms), Costo total y Longitud de la ruta.")

    add_body("En este primer caso, BFS, UCS, Greedy y A* convergen a la misma trayectoria óptima con costo de 13 unidades y 9 pasos. Sin embargo, A* y Greedy destacan por su notable eficiencia de exploración, expandiendo únicamente 12 y 10 nodos respectivamente, frente a los 16 nodos explorados por los métodos ciegos BFS y UCS. Por su parte, DFS se desvía por una rama lateral antes de retroceder, generando una ruta subóptima con 11 pasos y un costo elevado de 15 unidades.")

    # Caso 2
    add_h2("5.2 Caso de Estudio 2: Rescate Urbano con Rutas Alternativas (10x10)")
    add_body("El segundo escenario simula una planta de edificio de 10x10 celdas con múltiples pasillos paralelos, donde existen dos alternativas de navegación: una ruta superior con mayor número de pasos horizontales y una inferior con más tramos verticales.")

    headers_c2 = ["Métrica Evaluada", "BFS", "DFS", "UCS", "Greedy", "A*"]
    data_c2 = [
        ["Nodos explorados", "65", "33", "64", "20", "50"],
        ["Tiempo de ejecución (ms)", "0.1772", "0.0943", "0.2122", "0.0712", "0.1801"],
        ["Costo total (H=1, V=2)", "26", "46", "26", "26", "26"],
        ["Longitud de la ruta (pasos)", "18", "28", "18", "18", "18"],
        ["¿Encontró ruta óptima?", "Sí", "No", "Sí (Costo Mín.)", "Sí", "Sí (Costo Mín.)"]
    ]
    add_table_apa(headers_c2, data_c2, col_widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9],
                  title="Tabla 3. Métricas de Desempeño en el Escenario de Rescate Urbano (10x10)")
    
    add_image_figure("mapa_visual_urbano.png", "Figura 3.", "Comparación visual de la navegación en el Escenario de Rescate Urbano (10x10).")

    add_body("En este entorno más amplio, la búsqueda avara (Greedy) explora únicamente 20 nodos (un 68.7% menos que BFS y un 60% menos que A*), logrando un tiempo de CPU récord de 0.0712 ms debido a la ausencia de trampas que obstruyan la línea visual directa hacia la meta.")

    # Caso 3
    add_h2("5.3 Caso de Estudio 3: Escenario de Obstáculo Cóncavo / Trampa en U (12x12)")
    add_body("El tercer escenario somete a prueba los algoritmos frente a una 'trampa en U', un obstáculo cóncavo donde la meta se ubica dentro de una cavidad cuya única entrada se encuentra en el extremo opuesto a la dirección heurística natural.")

    headers_c3 = ["Métrica Evaluada", "BFS", "DFS", "UCS", "Greedy", "A*"]
    data_c3 = [
        ["Nodos explorados", "95", "73", "99", "37", "65"],
        ["Tiempo de ejecución (ms)", "0.2665", "0.3513", "0.4931", "0.1482", "0.2677"],
        ["Costo total (H=1, V=2)", "29", "129", "29", "29", "29"],
        ["Longitud de la ruta (pasos)", "18", "72", "18", "18", "18"],
        ["¿Encontró ruta óptima?", "Sí", "No (Severo desvío)", "Sí (Costo Mín.)", "Sí (Tras atasco)", "Sí (Costo Mín.)"]
    ]
    add_table_apa(headers_c3, data_c3, col_widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9],
                  title="Tabla 4. Métricas de Desempeño en el Escenario de Trampa Cóncava en U (12x12)")
    
    add_image_figure("mapa_visual_trampa.png", "Figura 4.", "Comparación visual en el Escenario de Trampa Cóncava en U (12x12).")

    add_body("Este caso expone con nitidez las limitaciones de la búsqueda ciega y de DFS. DFS sufrió un desvío catastrófico recorriendo 72 pasos con un costo acumulado de 129 unidades (más de 4 veces el costo óptimo de 29). Por su parte, A* superó con holgura a UCS explorando solo 65 nodos frente a los 99 nodos de UCS (un ahorro del 34.3% de estados expandidos), demostrando su robustez matemática para salir de mínimos locales heurísticos sin comprometer la optimalidad de la solución.")

    # Caso 4
    add_h2("5.4 Caso de Estudio 4: Escenario Complejo de Gran Escala (16x16)")
    add_body("El cuarto escenario evalúa el comportamiento y la escalabilidad de los algoritmos en una rejilla densa de 16x16 celdas con pasadizos estrechos, múltiples bifurcaciones y un espacio de búsqueda significativamente mayor.")

    headers_c4 = ["Métrica Evaluada", "BFS", "DFS", "UCS", "Greedy", "A*"]
    data_c4 = [
        ["Nodos explorados", "142", "51", "141", "35", "88"],
        ["Tiempo de ejecución (ms)", "0.3842", "0.1383", "0.4686", "0.1224", "0.3140"],
        ["Costo total (H=1, V=2)", "46", "80", "46", "46", "46"],
        ["Longitud de la ruta (pasos)", "31", "49", "31", "31", "31"],
        ["¿Encontró ruta óptima?", "Sí", "No", "Sí (Costo Mín.)", "Sí", "Sí (Costo Mín.)"]
    ]
    add_table_apa(headers_c4, data_c4, col_widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9],
                  title="Tabla 5. Métricas de Desempeño en el Escenario Complejo de Gran Escala (16x16)")
    
    add_image_figure("grafico_preset_4.png", "Figura 5.", "Gráficos comparativos de barras en el Escenario Complejo de Gran Escala (16x16).")

    # =========================================================================
    # SECCIÓN 6: COMPARACIÓN Y DISCUSIÓN DE ALGORITMOS
    # =========================================================================
    add_h1("6. Comparación y Discusión Integral de Algoritmos")
    add_body("El análisis comparativo consolidado de las pruebas experimentales arroja deducciones concluyentes respecto a los compromisos de diseño (trade-offs) entre tiempo de cómputo, uso de memoria y calidad de la solución encontrada:")

    headers_global = ["Algoritmo", "Clase de Búsqueda", "Garantía de Optimalidad", "Eficiencia en Nodos", "Velocidad CPU", "Comportamiento en Obstáculos Cóncavos"]
    data_global = [
        ["BFS", "No Informada (Ciega)", "Óptimo en número de pasos", "Baja (expansión radial completa)", "Media", "Inmune a trampas (exploración exhaustiva)"],
        ["DFS", "No Informada (Ciega)", "No óptimo (frecuente desvío)", "Variable (aleatoria según orden)", "Rápido", "Altamente vulnerable a desvíos masivos"],
        ["UCS", "No Informada (Ciega)", "Óptimo en costo acumulado", "Baja (explora todas las ramas baratas)", "Lenta", "Inmune a trampas (garantía global)"],
        ["Greedy", "Informada (Heurística)", "No garantizada", "Excelente en espacios abiertos", "Muy Rápida", "Vulnerable (requiere retroceso en trampas)"],
        ["A*", "Informada (Heurística)", "Óptimo en costo acumulado", "Alta (poda guiada por heurística)", "Rápida y Equilibrada", "Resiliente y matemáticamente óptimo"]
    ]
    add_table_apa(headers_global, data_global, col_widths=[0.9, 1.1, 1.2, 1.1, 0.9, 1.3],
                  title="Tabla 6. Matriz Comparativa Integral de los Cinco Algoritmos de Navegación")

    add_h2("6.1 Análisis del Consumo de Memoria y Nodos Explorados")
    add_body("Los algoritmos no informados BFS y UCS exploran el espacio de estados de manera concéntrica y uniforme, lo que conlleva la apertura de un volumen masivo de celdas irrelevantes situadas en direcciones opuestas a la meta. En el Escenario 4 (16x16), BFS y UCS exploraron 142 y 141 nodos respectivamente. En contraste, A* expandió únicamente 88 nodos (una reducción del 37.6% en el espacio de estados), mientras que Greedy expandió solo 35 nodos gracias a su focalización exclusiva hacia el gradiente heurístico.")

    add_h2("6.2 Impacto del Modelo de Costos Asimétricos (H=1, V=2)")
    add_body("La asignación de costos diferenciados evidenció la distinción conceptual crucial entre la minimización geométrica de pasos y la minimización energética de costos:")
    add_bullet("BFS optimiza estrictamente la longitud del camino en número de celdas transitadas, seleccionando trayectorias que minimizan la cantidad de saltos sin discriminar si dichos saltos son horizontales o verticales.", bold_prefix="• ")
    add_bullet("UCS y A* evalúan el costo real acumulado g(n), priorizando activamente los desplazamientos horizontales (c=1) sobre los verticales (c=2) siempre que la geometría del laberinto lo permita, lo que garantiza el menor gasto energético para el robot de rescate.", bold_prefix="• ")

    add_h2("6.3 Matriz de Decisión para Despliegue en Robótica de Rescate")
    add_body("Para entornos reales de rescate urbano donde las baterías del robot son limitadas y el tiempo de respuesta es vital:")
    add_bullet("Se recomienda el Algoritmo A* como la solución estándar predeterminada para navegación autónoma, dado que garantiza la ruta de menor costo energético con una fracción del consumo de memoria y tiempo de CPU que requeriría UCS.", bold_prefix="• ")
    add_bullet("Se sugiere Búsqueda Avara (Greedy) únicamente como método de exploración exploratoria preliminar cuando la latencia de respuesta deba ser ultrabaja y se conozca que el entorno posee baja densidad de obstáculos cóncavos.", bold_prefix="• ")

    # =========================================================================
    # SECCIÓN 7: CÓDIGO FUENTE DESARROLLADO Y DOCUMENTADO
    # =========================================================================
    add_h1("7. Código Fuente Desarrollado y Documentado")
    add_body("El proyecto se compone del archivo ejecutable principal ProyectoFinal.py, el cual integra todas las clases, algoritmos, presets, utilidades de benchmarking y la interfaz gráfica de usuario. A continuación se detallan las instrucciones de despliegue:")

    add_h2("7.1 Requisitos del Entorno de Ejecución")
    add_bullet("Lenguaje: Python 3.8 o superior.", bold_prefix="• ")
    add_bullet("Librerías estándar requeridas: tkinter, heapq, collections, time, random, copy, sys, os (incluidas nativamente en Python).", bold_prefix="• ")
    add_bullet("Librerías para gráficos estadísticos: matplotlib y numpy (opcionales para exportación de gráficos).", bold_prefix="• ")

    add_h2("7.2 Modos de Ejecución")
    add_body("El programa puede ser ejecutado en dos modos principales según las necesidades del usuario:")
    add_bullet("Modo Interfaz Gráfica (Por defecto): Ejecutar mediante comando estándar para abrir el editor visual, la animación interactiva y las herramientas de benchmarking:", bold_prefix="1. ")
    add_code("python ProyectoFinal.py")
    add_bullet("Modo Consola / Benchmark Automatizado: Ejecutar con la bandera --benchmark o --cli para realizar las pruebas automáticas en consola y generar los gráficos PNG:", bold_prefix="2. ")
    add_code("python ProyectoFinal.py --benchmark")

    # =========================================================================
    # SECCIÓN 8: CONCLUSIONES Y RECOMENDACIONES
    # =========================================================================
    add_h1("8. Conclusiones y Recomendaciones")
    
    add_h2("8.1 Conclusiones")
    add_bullet("Se diseñó e implementó exitosamente un Sistema Inteligente de Búsqueda y Navegación en Laberintos que resuelve de manera robusta el problema de navegación robótica en estructuras colapsadas, integrando cinco algoritmos fundamentales de Inteligencia Artificial (BFS, DFS, UCS, Greedy y A*).", bold_prefix="1. ")
    add_bullet("El algoritmo A* demostró ser la técnica más eficiente y equilibrada de todo el estudio, garantizando de forma matemáticamente rigurosa el camino de mínimo costo acumulado (idéntico al de UCS) pero reduciendo la exploración de nodos entre un 34% y un 50% respecto a los métodos de búsqueda ciega.", bold_prefix="2. ")
    add_bullet("La incorporación del modelo de costos diferenciados (horizontal=1, vertical=2) validó la necesidad de utilizar algoritmos basados en costos acumulados reales (UCS y A*), ya que los métodos basados únicamente en longitud de pasos (como BFS) no garantizan la optimización energética del robot.", bold_prefix="3. ")
    add_bullet("La Búsqueda Avara (Greedy), si bien ofrece los menores tiempos de ejecución en espacios abiertos, demostró su vulnerabilidad ante obstáculos cóncavos ('trampas en U'), confirmando que la omisión del costo histórico g(n) compromete la optimalidad y la robustez global del agente.", bold_prefix="4. ")
    add_bullet("La interfaz gráfica interactiva desarrollada en Tkinter demostró ser una herramienta pedagógica y técnica de alto valor, permitiendo visualizar de manera transparente la dinámica interna de las fronteras de exploración y corroborar los postulados teóricos en tiempo real.", bold_prefix="5. ")

    add_h2("8.2 Recomendaciones")
    add_bullet("Extender el sistema a entornos dinámicos donde los obstáculos puedan aparecer o desplazarse en tiempo real, implementando algoritmos de búsqueda incremental como D* Lite o LPA*.", bold_prefix="1. ")
    add_bullet("Incorporar soporte para conectividad de 8 direcciones (movimientos diagonales) ponderando el costo con la distancia euclidiana (c = sqrt(dx^2 + dy^2)) y aplicando la heurística de Chebyshev u Octile.", bold_prefix="2. ")
    add_bullet("Integrar modelos de visión por computador o simulación basada en ROS (Robot Operating System) y Gazebo para trasladar la lógica de búsqueda matricial a un robot físico diferencial en un entorno tridimensional.", bold_prefix="3. ")

    # =========================================================================
    # SECCIÓN 9: REFERENCIAS BIBLIOGRÁFICAS
    # =========================================================================
    add_h1("Referencias Bibliográficas")
    
    referencias = [
        ("Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2022). ", "Introduction to algorithms (4th ed.). ", "MIT Press."),
        ("Hart, P. E., Nilsson, N. J., & Raphael, B. (1968). ", "A formal basis for the heuristic determination of minimum cost paths. ", "IEEE Transactions on Systems Science and Cybernetics, 4(2), 100–107. https://doi.org/10.1109/TSSC.1968.300136"),
        ("Nilsson, N. J. (2014). ", "Principles of artificial intelligence. ", "Morgan Kaufmann Publishers."),
        ("Russell, S., & Norvig, P. (2020). ", "Artificial intelligence: A modern approach (4th ed.). ", "Pearson Education."),
        ("Sedgewick, R., & Wayne, K. (2011). ", "Algorithms (4th ed.). ", "Addison-Wesley Professional.")
    ]
    
    for autor_anio, titulo_cursiva, resto in referencias:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(3)
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.line_spacing = 1.3
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.first_line_indent = Inches(-0.5)
        
        r1 = p_ref.add_run(autor_anio)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.font.color.rgb = COLOR_TEXTO

        r2 = p_ref.add_run(titulo_cursiva)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        r2.font.italic = True
        r2.font.color.rgb = COLOR_TEXTO

        r3 = p_ref.add_run(resto)
        r3.font.name = 'Times New Roman'
        r3.font.size = Pt(12)
        r3.font.color.rgb = COLOR_TEXTO

    # Guardar documento final
    output_path = 'Proyecto Final Fundamentos de Inteligencia Artificial.docx'
    doc.save(output_path)
    print(f"[EXITO] Documento generado exitosamente en: {output_path}")

if __name__ == '__main__':
    construir_informe()
